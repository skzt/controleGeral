"""
Created on 16 de maio de 2016

@author: Pedro Vaz
"""
import os
from tkinter import Frame, Button, Entry, Toplevel, Label, messagebox

from docx import Document
from docx.shared import Pt
from win32com import client

from .Config import get_secret


class customBox(object):

    def __init__(self, parent, pedido):
        self.root = parent
        self.pedido = pedido
        self.top = Toplevel(self.root)
        self.top.wm_title("Usuário")
        self.pathLocation = get_secret("MAIN_PATH")

    def printBox(self):

        frame = Frame(self.top, borderwidth=4, relief='ridge')
        frame.pack(fill='both', expand=True)

        Label(frame, text="Nome").pack()

        self.entry = Entry(frame)
        self.entry.bind("<Return>", self.setNext)
        self.entry.focus()
        self.entry.pack(pady=4)

        imprimir = Button(frame)
        imprimir['text'] = 'Imprimir'
        imprimir['command'] = self.fillDocx
        imprimir.bind("<Return>", lambda _: self.fillDocx())
        imprimir.pack()

        cancel = Button(frame)
        cancel['text'] = 'Cancelar'
        cancel['command'] = self.top.destroy
        cancel.bind("<Return>", lambda _: self.top.destroy())
        cancel.pack(padx=4, pady=4)

    def fillDocx(self):

        # Abre o documento base.
        _path = os.path.abspath(os.path.join(self.pathLocation, "toPrint.docx"))
        doc = Document(_path)
        # Carrega a tabela na memoria
        table = doc.tables[0]

        # Configura a fonte e tamanho da letra a ser escrita
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        for row in table.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs:
                    paragrafo = doc.styles['Normal']
        # CELULA     CONTEUDO
        #     0,5        CI
        #     1,1        De:
        #     1,4        Para:
        #     2,1        Data:
        #     3,0        Produtos:
        #     5,5        CI
        #     6,1        De:
        #     6,4        Para:
        #     7,1        Data:
        #     8,0        Produtos:

        table.cell(0, 5).text = self.pedido[0]
        table.cell(1, 1).text = self.entry.get() + " - CPD"
        table.cell(1, 4).text = self.pedido[2] + " - " + self.pedido[3]
        table.cell(2, 1).text = self.pedido[4]
        table.cell(3, 0).text = self.pedido[5] + " - " + self.pedido[6]

        table.cell(5, 5).text = self.pedido[0]
        table.cell(6, 1).text = self.entry.get() + " - CPD"
        table.cell(6, 4).text = self.pedido[2] + " - " + self.pedido[3]
        table.cell(7, 1).text = self.pedido[4]
        table.cell(8, 0).text = self.pedido[5] + " - " + self.pedido[6]

        save = os.path.join(self.pathLocation, "Aberto", self.pedido[3], f"{self.pedido[0]}.docx")
        doc.save(os.path.abspath(save))
        imprimir = messagebox.askyesno(title="Imprimir C.I",
                                       message="Deseja Imprimir o Controle Interno para este pedido?",
                                       parent=self.top)
        if imprimir is True:
            self.printWordDocument(save)

        self.top.destroy()

    def printWordDocument(self, filename):
        word = client.Dispatch("Word.Application")
        word.Documents.Open(filename)
        word.ActiveDocument.PrintOut()
        word.ActiveDocument.Close()
        word.Quit()

    def setNext(self, event):
        event.widget.tk_focusNext().focus()
        return ("break")
