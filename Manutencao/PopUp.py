"""
Created on 25 de jun de 2018

@author: GYN-CPD-PEDRO
"""

import tkinter as tk
import tkinter.ttk as ttk

from docx import Document


class popUP(tk.Toplevel):
    def __init__(self, parent, instance, *arg, **kwarg):
        tk.Toplevel.__init__(self, parent, *arg, **kwarg)

        self.__instance = instance
        self.__selecionado = None
        self.frame = tk.Frame(self)
        self.frame.grid(row=0, column=0)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self.sair)

        exitButton = tk.Button(self)
        exitButton['text'] = "Sair"
        exitButton['command'] = self.sair
        exitButton.grid(row=1, column=0, sticky='e', pady=(5, 20), padx=(0, 20))

    def montarTree(self, columns, dados):

        self.focus()
        self.bind("<Escape>", lambda _: self.sair())

        selectButton = tk.Button(self)
        selectButton['text'] = "Carregar Dados"
        selectButton['command'] = lambda: self.onSelection(obs=dados)
        selectButton.grid(row=1, column=0, sticky='w', pady=(5, 20), padx=(20, 0))

        self.tree = ttk.Treeview(self.frame)
        self.tree.bind("<Double-1>", lambda _: self.onSelection(obs=dados))

        if self.__instance.upper() == 'CC':

            self.tree['columns'] = ('0', '1', '2', '3', '4', '5', '6')
            self.tree.column('#0', width=10)
            self.tree.column('6', width=10)

            self.tree.column('0', width=125)  # Num Chamado
            self.tree.heading('0',
                              text=columns[0],
                              command=lambda: self.sortCol('0', False))

            self.tree.column('1', width=100)  # Num Serie
            self.tree.heading('1',
                              text=columns[1],
                              command=lambda: self.sortCol('1', False))

            self.tree.column('2', width=75)  # Data
            self.tree.heading('2',
                              text=columns[2],
                              command=lambda: self.sortCol('2', False))

            self.tree.column('3', width=70)  # Hora
            self.tree.heading('3',
                              text=columns[3],
                              command=lambda: self.sortCol('3', False))

            self.tree.column('4', width=50)  # Filial
            self.tree.heading('4',
                              text=columns[4],
                              command=lambda: self.sortCol('4', False))

            self.tree.column('5', width=100)  # Responsavel
            self.tree.heading('5',
                              text=columns[5],
                              command=lambda: self.sortCol('5', False))

            for line in dados:
                self.tree.insert('', 'end', values=line[:6])

            dicDados = {}
            index = 0

            for iid in self.tree.get_children(''):
                dicDados[iid] = dados[index][6]
                index += 1

            dados = dicDados
            del dicDados

        elif self.__instance.upper() == 'CV':
            self.tree['columns'] = ('0', '1', '2', '3', '4')
            self.tree.column('#0', width=10)
            self.tree.column('4', width=10)

            self.tree.column('0', width=125)  # Num Chamado
            self.tree.heading('0',
                              text=columns[0],
                              command=lambda: self.sortCol('0', False))

            self.tree.column('1', width=75)  # Data
            self.tree.heading('1',
                              text=columns[1],
                              command=lambda: self.sortCol('1', False))

            self.tree.column('2', width=70)  # Hora
            self.tree.heading('2',
                              text=columns[2],
                              command=lambda: self.sortCol('2', False))

            self.tree.column('3', width=100)  # Resolvido
            self.tree.heading('3',
                              text=columns[3],
                              command=lambda: self.sortCol('3', False))

            for line in dados:
                tmp = line[:3]
                tmp.append(line[4])
                self.tree.insert('', 'end', values=tmp)
                del tmp

            dicDados = {}
            index = 0
            for iid in self.tree.get_children(''):
                dicDados[iid] = dados[index][3]
                index += 1

            dados = dicDados


        elif self.__instance == 'CP':
            self.tree['columns'] = ('0', '1', '2', '3', '4', '5')
            self.tree.column('#0', width=10)
            self.tree.column('5', width=10)

            self.tree.column('0', width=125)  # Num Chamado
            self.tree.heading('0',
                              text=columns[0],
                              command=lambda: self.sortCol('0', False))

            self.tree.column('1', width=75)  # NF Entrada
            self.tree.heading('1',
                              text=columns[1],
                              command=lambda: self.sortCol('1', False))

            self.tree.column('2', width=70)  # Data Envio
            self.tree.heading('2',
                              text=columns[2],
                              command=lambda: self.sortCol('2', False))

            self.tree.column('3', width=100)  # NF Saida
            self.tree.heading('3',
                              text=columns[3],
                              command=lambda: self.sortCol('3', False))

            self.tree.column('4', width=110)  # Data Recebimento
            self.tree.heading('4',
                              text=columns[4],
                              command=lambda: self.sortCol('4', False))
            for line in dados:
                tmp = line
                self.tree.insert('', 'end', values=tmp)
                del tmp

            dicDados = {}
            index = 0
            for iid in self.tree.get_children(''):
                dicDados[iid] = dados[index][3]
                index += 1

            dados = dicDados

        elif self.__instance == 'VC':
            self.tree.unbind("<Double-1>")

            path = dados[0]
            dados = dados[1]

            tmp = []
            for ci in dados:
                doc = Document(path + ci)
                tmp.append([ci, doc.tables[0].cell(1, 4).text, doc.tables[0].cell(2, 1).text])

            dados = tmp
            del tmp

            self.tree['columns'] = ('0', '1', '2', '3')

            self.tree.column('#0', width=10)
            self.tree.column('3', width=10)

            self.tree.column('0', width=125)  # CI
            self.tree.heading('0',
                              text=columns[0],
                              command=lambda: self.sortCol('0', False))

            self.tree.column('1', width=125)  # SOLICITANTE
            self.tree.heading('1',
                              text=columns[1],
                              command=lambda: self.sortCol('1', False))

            self.tree.column('2', width=125)  # DATA DE ENVIO
            self.tree.heading('2',
                              text=columns[2],
                              command=lambda: self.sortCol('2', False))

            for line in dados:
                self.tree.insert('', 'end', values=line)

            selectButton.grid_forget()

        self.tree.grid(row=0, column=0, pady=(20, 5), padx=(20, 20))
        self.wait_window(self)

    def montarTreeColunaUnica(self, dados, column, widget):
        self.focus()

        self.bind("<Escape>", lambda _: self.sair())

        selectButton = tk.Button(self)
        selectButton['text'] = "Carregar Dados"
        selectButton['command'] = lambda dados=dados: self.onSelection(dados, True, widget)
        selectButton.grid(row=1, column=0, sticky='w', pady=(5, 20), padx=(20, 0))

        self.tree = ttk.Treeview(self.frame)
        self.tree.bind("<Double-1>", lambda _, dados=dados: self.onSelection(dados, True, widget))

        self.tree['column'] = ('0', '1')
        self.tree.column('#0', width=10)
        self.tree.column('1', width=10)

        self.tree.column('0', width=150)
        self.tree.heading('0',
                          text=column,
                          command=lambda: self.sortCol('0', False))

        for line in dados:
            self.tree.insert('', 'end', values=line[0])

        self.tree.grid(row=0, column=0, pady=(20, 5), padx=(20, 20))
        self.wait_window(self)

    def sortCol(self, col, reverse):
        line = [(self.tree.set(iid, col), iid) for iid in self.tree.get_children('')]

        line.sort(reverse=reverse)

        for index, (_, iid) in enumerate(line):
            self.tree.move(iid, '', index)

        self.tree.heading(col, command=lambda: self.sortCol(col, not reverse))

    def onSelection(self, obs=None, chamado=None, widget=None):
        iid = self.tree.selection()[0]
        dados = self.tree.set(iid)

        if len(dados) == 1:
            if chamado:
                self.master.entradas[widget].focus()
            self.master.entradas[widget].delete(0, 'end')
            self.master.entradas[widget].insert('end', dados['0'])

        else:
            if self.__instance == 'CC':
                self.master.new = False
            index = 0
            for w in self.master.entradas:

                if isinstance(w, tk.Text):
                    w.delete(0.0, 'end')
                    w.insert('end', obs[iid])
                else:
                    w.delete(0, 'end')
                    w.insert('end', dados[str(index)])
                if self.__instance == 'CV' and index == 3:
                    tmp = 1 if dados[str(index)] == 'Sim' else 0
                    self.master.variaveisVisitas[3].set(tmp)
                index += 1

        self.destroy()

    def sair(self):
        self.master.err = True
        self.destroy()
